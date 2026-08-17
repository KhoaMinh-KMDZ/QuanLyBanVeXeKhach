from __future__ import annotations

import os
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = Path(r"C:\Users\tuant\OneDrive\Desktop\Doc1.docx")
OUT_DIR = ROOT / "BaoCao"
MEDIA_DIR = OUT_DIR / "anh_giao_dien"
OUT_FILE = OUT_DIR / "BaoCao_DeTai_QuanLyBanVeXeKhach.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill: str = "F2F4F7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_para(doc: Document, text: str = "", bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def extract_images() -> list[Path]:
    MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    images: list[Path] = []
    if not SOURCE_DOCX.exists():
        return images
    with zipfile.ZipFile(SOURCE_DOCX) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                out = MEDIA_DIR / Path(name).name
                out.write_bytes(zf.read(name))
                images.append(out)
    return sorted(images)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    for text, size, bold in [
        ("TRƯỜNG CAO ĐẲNG/ĐẠI HỌC: ................................", 13, True),
        ("KHOA: CÔNG NGHỆ THÔNG TIN", 13, True),
        ("BÁO CÁO ĐỀ TÀI", 20, True),
        ("MÔN: CÔNG NGHỆ .NET", 15, True),
        ("Đề tài: XÂY DỰNG PHẦN MỀM QUẢN LÝ BÁN VÉ XE KHÁCH", 17, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        if "Đề tài" in text or text == "BÁO CÁO ĐỀ TÀI":
            r.font.color.rgb = RGBColor(31, 77, 120)

    doc.add_paragraph()
    info = [
        ("Lớp", "................................"),
        ("Nhóm thực hiện", "Nhóm 5"),
        ("Giảng viên hướng dẫn", "................................"),
        ("Thành viên 1", "................................"),
        ("Thành viên 2", "................................"),
        ("Thành viên 3", "................................"),
        ("Thành viên 4", "................................"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.2)
    for i, (k, v) in enumerate(info):
        cells = table.rows[0].cells if i == 0 else table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    style_table(table)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.add_run("TP. Hồ Chí Minh, tháng 06 năm 2026").italic = True
    doc.add_page_break()


def add_group_front_matter(doc: Document) -> None:
    add_heading(doc, "NHẬN XÉT CỦA GIẢNG VIÊN", 1)
    for _ in range(12):
        add_para(doc, "................................................................................................................")
    doc.add_page_break()

    add_heading(doc, "BẢNG PHÂN CÔNG CÔNG VIỆC", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["STT", "Họ tên", "Công việc", "Mức độ hoàn thành"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    tasks = [
        ("1", "Thành viên 1", "Phân tích yêu cầu, thiết kế CSDL", "100%"),
        ("2", "Thành viên 2", "Xây dựng giao diện WPF", "100%"),
        ("3", "Thành viên 3", "Xử lý nghiệp vụ bán vé, chuyến xe", "100%"),
        ("4", "Thành viên 4", "Báo cáo, kiểm thử, hoàn thiện", "100%"),
    ]
    for row_data in tasks:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "NHẬT KÝ THỰC HIỆN NHÓM", 1)
    log = doc.add_table(rows=1, cols=4)
    log.style = "Table Grid"
    for i, h in enumerate(["Tuần", "Nội dung thực hiện", "Người phụ trách", "Kết quả"]):
        set_cell_text(log.rows[0].cells[i], h, True)
    for row in [
        ("Tuần 1", "Khảo sát đề tài, xác định chức năng và dữ liệu cần quản lý.", "Cả nhóm", "Hoàn thành yêu cầu tổng quát."),
        ("Tuần 2", "Thiết kế database, bảng dữ liệu, khóa chính/khóa ngoại, stored procedure.", "Nhóm CSDL", "Có script SQL hoàn chỉnh."),
        ("Tuần 3", "Xây dựng giao diện đăng nhập, trang chủ, bán vé, danh mục.", "Nhóm giao diện", "Hoàn thành các màn hình chính."),
        ("Tuần 4", "Lập trình ViewModel/Repository, kiểm thử đăng nhập admin và nghiệp vụ bán vé.", "Cả nhóm", "Chạy được chức năng chính."),
        ("Tuần 5", "Viết báo cáo, rà soát lỗi và hoàn thiện nộp bài.", "Cả nhóm", "Hoàn thiện báo cáo đề tài."),
    ]:
        cells = log.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(log)
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    for item in [
        "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI",
        "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 3. THIẾT KẾ CƠ SỞ DỮ LIỆU",
        "CHƯƠNG 4. THIẾT KẾ VÀ CÀI ĐẶT CHƯƠNG TRÌNH",
        "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ",
        "KẾT LUẬN",
        "PHỤ LỤC",
    ]:
        add_para(doc, item)
    doc.add_page_break()


def add_main_content(doc: Document, images: list[Path]) -> None:
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_para(doc, "Hoạt động bán vé xe khách cần xử lý nhiều thông tin như tuyến xe, chuyến xe, xe, tài xế, khách hàng, ghế ngồi và doanh thu. Nếu quản lý thủ công bằng sổ sách hoặc file rời, việc tra cứu vé, kiểm soát ghế trống và tổng hợp doanh thu dễ sai sót. Vì vậy, nhóm chọn đề tài xây dựng phần mềm quản lý bán vé xe khách bằng công nghệ .NET nhằm mô phỏng một hệ thống quản lý thực tế.")
    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    for item in [
        "Xây dựng ứng dụng desktop bằng WPF trên nền .NET 8.0.",
        "Quản lý đăng nhập, phân quyền theo vai trò nhân viên.",
        "Quản lý danh mục bến xe, tuyến xe, xe, khách hàng, nhân viên.",
        "Quản lý chuyến xe, sơ đồ ghế, bán vé, hủy vé và lịch sử giao dịch.",
        "Thống kê doanh thu theo ngày, theo tuyến và hỗ trợ xuất Excel.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3. Phạm vi sử dụng", 2)
    add_para(doc, "Hệ thống phù hợp với bài tập học phần Công nghệ .NET, mô phỏng nghiệp vụ tại nhà xe/quầy bán vé. Dữ liệu lưu trên SQL Server, giao diện chạy trên Windows bằng WPF.")

    add_heading(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "2.1. Tác nhân sử dụng", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Tác nhân", "Quyền chính", "Mô tả"]):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in [
        ("Admin", "Toàn quyền", "Quản lý bán vé, chuyến xe, danh mục, báo cáo và nhân viên."),
        ("Nhân viên bán vé", "Bán vé", "Chọn chuyến, chọn ghế, nhập khách hàng, thanh toán, tra cứu/hủy vé."),
        ("Quản lý điều hành", "Điều hành và báo cáo", "Quản lý chuyến, tuyến, danh mục và xem doanh thu."),
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    add_heading(doc, "2.2. Các chức năng chính", 2)
    for item in [
        "Đăng nhập: kiểm tra tài khoản, mật khẩu đã băm SHA256 và trạng thái tài khoản.",
        "Trang chủ: hiển thị số vé bán hôm nay, doanh thu hôm nay, chuyến sắp chạy và chuyến đang chạy.",
        "Bán vé: chọn chuyến xe, xem sơ đồ ghế, nhập số điện thoại/họ tên khách, chọn hình thức thanh toán và lưu vé.",
        "Quản lý vé: tra cứu vé theo số điện thoại, xem thông tin vé và hủy vé đã chọn.",
        "Quản lý chuyến xe: thêm/sửa/xóa chuyến, cập nhật trạng thái sắp chạy, đang chạy, hoàn thành, đã hủy.",
        "Quản lý danh mục: bến xe, tuyến xe, xe, khách hàng.",
        "Quản trị: quản lý nhân viên, vai trò, khóa/mở khóa tài khoản và đổi mật khẩu.",
        "Báo cáo: xem doanh thu theo khoảng ngày, thống kê tuyến có doanh thu cao và xuất file Excel.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "CHƯƠNG 3. THIẾT KẾ CƠ SỞ DỮ LIỆU", 1)
    add_para(doc, "Cơ sở dữ liệu sử dụng SQL Server với tên QuanLyBanVeXeKhach. Script tạo database có các bảng chính, ràng buộc khóa ngoại, trigger, function, stored procedure và dữ liệu mẫu phục vụ kiểm thử.")
    add_heading(doc, "3.1. Danh sách bảng dữ liệu", 2)
    tables = [
        ("VaiTro", "Lưu vai trò và quyền: bán vé, quản lý chuyến, danh mục, báo cáo, nhân viên."),
        ("NhanVien", "Lưu nhân viên, tài khoản, mật khẩu băm SHA256 và vai trò."),
        ("KhachHang", "Lưu họ tên, số điện thoại, email, địa chỉ, ghi chú khách hàng."),
        ("LoaiXe", "Lưu loại xe và số ghế."),
        ("Xe", "Lưu biển số, loại xe, năm sản xuất, trạng thái."),
        ("BenXe", "Lưu tên bến xe, tỉnh thành, địa chỉ, trạng thái."),
        ("TuyenXe", "Lưu tuyến, bến đi, bến đến, khoảng cách, thời gian di chuyển."),
        ("ChuyenXe", "Lưu lịch trình xuất bến, xe, tài xế, giá vé, trạng thái."),
        ("VeXe", "Lưu vé, khách hàng, nhân viên bán, số ghế, giá vé, trạng thái."),
        ("NhatKyHeThong", "Lưu nhật ký bán vé, hủy vé, cập nhật vé từ trigger."),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["STT", "Bảng", "Ý nghĩa"]):
        set_cell_text(t.rows[0].cells[i], h, True)
    for idx, (name, desc) in enumerate(tables, 1):
        cells = t.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = name
        cells[2].text = desc
    style_table(t)

    add_heading(doc, "3.2. Ràng buộc và xử lý trong SQL", 2)
    for item in [
        "Khóa ngoại liên kết nhân viên - vai trò, xe - loại xe, tuyến - bến xe, chuyến - tuyến/xe/tài xế, vé - chuyến/khách hàng/nhân viên.",
        "Unique index UQ_VeXe_Ghe_Active đảm bảo một ghế trong một chuyến chỉ có một vé đang đặt hoặc đã thanh toán.",
        "Trigger trg_NhatKyBanVe và trg_NhatKyCapNhatVe tự động ghi nhật ký khi bán hoặc cập nhật trạng thái vé.",
        "Trigger trg_NganXoaKhachHang ngăn xóa khách hàng đã có lịch sử mua vé.",
        "Function fn_DoanhThuTheoNgay và fn_SoGheTrong hỗ trợ thống kê doanh thu, số ghế trống.",
        "Stored procedure sp_ThanhToanVeXe dùng transaction và khóa UPDLOCK/SERIALIZABLE để hạn chế tranh chấp ghế.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ VÀ CÀI ĐẶT CHƯƠNG TRÌNH", 1)
    add_heading(doc, "4.1. Công nghệ sử dụng", 2)
    for item in [
        "Ngôn ngữ: C#.",
        "Framework: .NET 8.0 Windows.",
        "Giao diện: WPF kết hợp MaterialDesignThemes.",
        "Cơ sở dữ liệu: SQL Server, thư viện Microsoft.Data.SqlClient.",
        "Xuất báo cáo: ClosedXML để tạo file Excel.",
        "Kiến trúc: MVVM gồm Models, Views, ViewModels, Repositories, Helpers, Services.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2. Cấu trúc project", 2)
    project_table = doc.add_table(rows=1, cols=2)
    project_table.style = "Table Grid"
    for i, h in enumerate(["Thư mục/File", "Vai trò"]):
        set_cell_text(project_table.rows[0].cells[i], h, True)
    for row in [
        ("Models", "Định nghĩa đối tượng dữ liệu: NhanVien, KhachHang, Xe, TuyenXe, ChuyenXe, VeXe..."),
        ("Views", "Các màn hình WPF XAML: đăng nhập, trang chủ, bán vé, quản lý danh mục, báo cáo..."),
        ("ViewModels", "Xử lý binding, command và nghiệp vụ phía giao diện theo MVVM."),
        ("Repositories", "Truy vấn SQL Server, gọi stored procedure và ánh xạ dữ liệu."),
        ("Helpers", "Hỗ trợ kết nối CSDL, băm mật khẩu, session đăng nhập, sơ đồ ghế, thông báo."),
        ("Services", "Xử lý xuất báo cáo Excel."),
        ("CSDL_QuanLyBanVeXeKhach_Full.sql", "Script tạo database, bảng, trigger, function, procedure và dữ liệu mẫu."),
    ]:
        cells = project_table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(project_table)

    add_heading(doc, "4.3. Luồng đăng nhập", 2)
    for step in [
        "Người dùng nhập tài khoản và mật khẩu trên màn hình đăng nhập.",
        "ViewModel kiểm tra dữ liệu rỗng, sau đó băm mật khẩu bằng SHA256.",
        "NhanVienRepository gọi stored procedure sp_DangNhap để kiểm tra tài khoản, mật khẩu hash và trạng thái đang hoạt động.",
        "Nếu hợp lệ, SessionManager lưu thông tin người dùng và quyền, sau đó mở MainView.",
        "Menu bên trái hiển thị theo quyền của vai trò hiện tại.",
    ]:
        add_number(doc, step)
    add_para(doc, "Tài khoản kiểm thử: admin / 123456. Trong SQL, mật khẩu 123456 được lưu dưới dạng SHA256: 8d969eef6ecad3c29a3a629280e686cf0c3f5d5a86aff3ca12020c923adc6c92.")

    add_heading(doc, "4.4. Luồng bán vé", 2)
    for step in [
        "Nhân viên chọn chuyến xe còn hoạt động.",
        "Hệ thống tải thông tin tuyến, xe, giá vé, số ghế trống và danh sách ghế đã bán.",
        "Sơ đồ ghế được tạo theo loại xe: limousine 11/16 chỗ, xe 29 chỗ, xe 45 chỗ hai tầng.",
        "Nhân viên chọn ghế trống, nhập số điện thoại và họ tên khách hàng.",
        "Khi xác nhận thanh toán, hệ thống gọi sp_ThanhToanVeXe để tạo hoặc cập nhật khách hàng và thêm vé trong cùng transaction.",
        "Trigger tự ghi nhật ký bán vé vào bảng NhatKyHeThong.",
    ]:
        add_number(doc, step)

    add_heading(doc, "4.5. Giao diện chương trình", 2)
    captions = [
        "Màn hình đăng nhập hệ thống.",
        "Giao diện chính và menu điều hướng theo phân quyền.",
        "Trang chủ thống kê nhanh vé, doanh thu và chuyến xe.",
        "Màn hình bán vé và chọn chuyến xe.",
        "Sơ đồ ghế phục vụ chọn ghế.",
        "Hộp thoại thanh toán vé.",
        "Màn hình quản lý chuyến xe.",
        "Màn hình quản lý tuyến xe.",
        "Màn hình quản lý bến xe.",
        "Màn hình quản lý danh mục xe.",
        "Màn hình quản lý nhân viên và phân quyền.",
        "Màn hình báo cáo doanh thu.",
    ]
    for idx, image in enumerate(images[:12], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(image), width=Inches(5.8))
        except Exception:
            p.add_run(f"[Không thể chèn ảnh: {image.name}]")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Hình {idx}. {captions[idx - 1] if idx <= len(captions) else 'Giao diện chương trình.'}")
        run.italic = True
        run.font.size = Pt(10)

    add_heading(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "5.1. Kịch bản kiểm thử", 2)
    test_table = doc.add_table(rows=1, cols=4)
    test_table.style = "Table Grid"
    for i, h in enumerate(["STT", "Chức năng", "Dữ liệu kiểm thử", "Kết quả mong đợi"]):
        set_cell_text(test_table.rows[0].cells[i], h, True)
    tests = [
        ("1", "Đăng nhập", "admin / 123456", "Đăng nhập thành công, vào màn hình chính với quyền Admin."),
        ("2", "Đăng nhập sai", "admin / mật khẩu sai", "Thông báo tài khoản hoặc mật khẩu không đúng."),
        ("3", "Bán vé", "Chọn chuyến, chọn ghế trống, nhập khách hàng", "Tạo vé, cập nhật ghế đã đặt, ghi nhật ký."),
        ("4", "Trùng ghế", "Hai thao tác chọn cùng một ghế", "SQL chặn bằng unique index/transaction, không tạo vé trùng."),
        ("5", "Hủy vé", "Tra cứu theo SĐT và hủy vé", "Trạng thái vé chuyển sang đã hủy, ghế được giải phóng."),
        ("6", "Báo cáo", "Chọn khoảng ngày", "Hiển thị số vé, doanh thu và xuất được Excel."),
    ]
    for row in tests:
        cells = test_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(test_table)

    add_heading(doc, "5.2. Đánh giá kết quả", 2)
    for item in [
        "Ứng dụng có đầy đủ các màn hình chính của nghiệp vụ quản lý bán vé xe khách.",
        "Cấu trúc MVVM giúp tách giao diện, xử lý nghiệp vụ và truy vấn dữ liệu.",
        "CSDL có ràng buộc, trigger, function và stored procedure hỗ trợ nghiệp vụ quan trọng.",
        "Phần bán vé có xử lý chống trùng ghế ở tầng CSDL.",
        "Giao diện sử dụng Material Design nên trực quan, dễ thao tác.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3. Hạn chế", 2)
    for item in [
        "Chưa có chức năng in vé chuyên nghiệp theo mẫu hóa đơn đầy đủ.",
        "Chưa có phân quyền chi tiết đến từng nút thao tác nhỏ.",
        "Chưa tích hợp thanh toán online thật, phần chuyển khoản mới ở mức mô phỏng.",
        "Chưa có module sao lưu/phục hồi tự động trên giao diện.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "KẾT LUẬN", 1)
    add_para(doc, "Đề tài đã xây dựng được phần mềm quản lý bán vé xe khách bằng C# WPF và SQL Server. Hệ thống đáp ứng các nghiệp vụ cơ bản như đăng nhập, phân quyền, quản lý danh mục, quản lý chuyến xe, bán vé theo sơ đồ ghế, hủy vé, thống kê doanh thu và xuất báo cáo Excel. Qua quá trình thực hiện, nhóm củng cố kiến thức về .NET, WPF, mô hình MVVM, kết nối SQL Server và thiết kế cơ sở dữ liệu có ràng buộc nghiệp vụ.")
    add_para(doc, "Hướng phát triển tiếp theo là bổ sung in vé chuẩn, thanh toán điện tử, phân quyền chi tiết hơn, sao lưu dữ liệu tự động và triển khai hệ thống cho nhiều máy trạm trong cùng mạng nội bộ.")

    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "A. Thông tin chạy chương trình", 2)
    for item in [
        "Mở project bằng Visual Studio 2022.",
        "Cài SQL Server/SQL Server Express và chạy file CSDL_QuanLyBanVeXeKhach_Full.sql.",
        "Chạy ứng dụng WPF, đăng nhập bằng tài khoản admin, mật khẩu 123456.",
        "Nếu cần đổi server SQL, kiểm tra cấu hình trong Helpers/DatabaseHelper.cs hoặc file dbconfig.txt ở thư mục build.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "B. Tài liệu tham khảo", 2)
    for item in [
        "Microsoft Docs - WPF và .NET.",
        "Microsoft Docs - SQL Server, T-SQL, Stored Procedure, Trigger.",
        "Tài liệu học phần Công nghệ .NET.",
        "Mã nguồn project QuanLyBanVeXeKhach của nhóm.",
    ]:
        add_bullet(doc, item)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Báo cáo đề tài - Quản lý bán vé xe khách"


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    images = extract_images()
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_group_front_matter(doc)
    add_main_content(doc, images)
    add_footer(doc)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
